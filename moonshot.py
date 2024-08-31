# -*- coding: utf-8 -*-
from pathlib import Path
from openai import OpenAI
from docx import Document
import os

client = OpenAI(
    api_key="sk-gyqGihWcBDXjapcnDnOuDeLDDTgQ2oNmM5DUBmd5urZhwpl9",  # 在这里将 MOONSHOT_API_KEY 替换为你从 Kimi 开放平台申请的 API Key
    base_url="https://api.moonshot.cn/v1",
)


def splitDoc(doc_file):
    # 加载文档
    input_doc = Document(doc_file)

    # 将文档内容读取为一个字符串
    full_text = []
    for paragraph in input_doc.paragraphs:
        full_text.append(paragraph.text)
    full_text = "\n".join(full_text)

    # 根据 "https://www.douyin.com/video/" 进行切分
    sections = full_text.split("https://www.douyin.com/video/")

    # 创建保存文件的文件夹
    output_dir = 'output_documents'
    os.makedirs(output_dir, exist_ok=True)

    # 遍历切分后的内容并保存为独立的Word文档
    for i, section in enumerate(sections):
        if section.strip():  # 忽略空白部分
            doc = Document()
            p = doc.add_paragraph("https://www.douyin.com/video/" + section.strip())

            # 确保文本不加粗
            for run in p.runs:
                run.bold = False  # 禁用加粗

            output_path = os.path.join(output_dir, f'document_{i + 1}.docx')
            doc.save(output_path)

    print(f"文档拆分完成，共生成 {len(sections) - 1} 个文件，保存在目录: {output_dir}")


def loopParse():
    # 定义要检查的文件夹路径
    output_dir = 'output_documents'

    # 判断文件夹是否存在
    if os.path.exists(output_dir) and os.path.isdir(output_dir):
        print(f"'{output_dir}' 文件夹存在，正在遍历其中的文件...")

        # 遍历文件夹中的文件
        for filename in os.listdir(output_dir):
            if filename.endswith('.docx') and not filename.startswith('~$'):  # 检查是否为Word文档
                moonparse(output_dir , filename)
                print(output_dir + "/" + filename)
    else:
        print(f"'{output_dir}' 文件夹不存在。")

def moonparse(output_dir, filename):
    # moonshot.pdf 是一个示例文件, 我们支持文本文件和图片文件，对于图片文件，我们提供了 OCR 的能力
    # 上传文件时，我们可以直接使用 openai 库的文件上传 API，使用标准库 pathlib 中的 Path 构造文件
    # 对象，并将其传入 file 参数即可，同时将 purpose 参数设置为 file-extract；注意，目前文件上传
    # 接口仅支持 file-extract 一种 purpose 值。
    file_object = client.files.create(file=Path(output_dir + "/" + filename), purpose="file-extract")

    # 获取结果
    # file_content = client.files.retrieve_content(file_id=file_object.id)
    # 注意，某些旧版本示例中的 retrieve_content API 在最新版本标记了 warning, 可以用下面这行代替
    # （如果使用旧版本的 SDK，可以继续延用 retrieve_content API）
    file_content = client.files.content(file_id=file_object.id).text

    # 把文件内容通过系统提示词 system prompt 放进请求中
    messages = [
        {
            "role": "system",
            "content": "你是 Kimi，由 Moonshot AI 提供的人工智能助手，你更擅长中文标点符号。你会为用户提供安全，有帮助，准确的回答。同时，你会拒绝一切涉及恐怖主义，种族歧视，黄色暴力等问题的回答。Moonshot AI 为专有名词，不可翻译成其他语言。"

        },
        {
            "role": "system",
            "content": file_content,  # <-- 这里，我们将抽取后的文件内容（注意是文件内容，而不是文件 ID）放置在请求中
        },
        {"role": "user", "content":
                                    "仅对文档中的“## 视频ASR文本：”部分进行标点符号的补全"
                                    "3. 补全标点符号后，确保语句连贯、意义明确"},
    ]

    # 然后调用 chat-completion, 获取 Kimi 的回答
    completion = client.chat.completions.create(
        model="moonshot-v1-128k",
        messages=messages,
        temperature=0.3,
    )

    result = str(completion.choices[0].message)
    result = (result.replace("ChatCompletionMessage(content='您好，","")
                      .replace("（文档结束）', refusal=None, role='assistant', function_call=None, tool_calls=None)","")
                      .replace("', refusal=None, role='assistant', function_call=None, tool_calls=None)","")
                      .replace("refusal=None, role='assistant', function_call=None, tool_calls=None)",""))

    # 加载文档document_modified_2
    input_doc = Document(output_dir + "/" + filename)

    # 遍历文档中的段落
    for paragraph in input_doc.paragraphs:
        if "## 视频ASR文本:" in paragraph.text:
            # 分割文本以保留 "## 视频ASR文本:" 前的内容
            parts = paragraph.text.split("## 视频ASR文本:")
            # 重新组合并替换后面的内容为 "暂无文档说明"
            paragraph.text = parts[0]  + result.replace("ChatCompletionMessage(content='## ","")

    # 保存修改后的文档
    output_path = 'output_documents_new/' + filename
    input_doc.save(output_path)


def write_to_markdown(text_content, md_filename='example.md'):
    """
    将文本内容写入指定的Markdown文件。

    参数:
    text_content (str): 要写入的文本内容。
    md_filename (str): Markdown文件的名称，默认为'example.md'。
    """
    try:
        # 打开文件并写入内容
        with open(md_filename, 'w', encoding='utf-8') as file:
            file.write(text_content)
        print(f"内容已写入到 {md_filename}")
    except Exception as e:
        print(f"写入文件时发生错误: {e}")



# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    splitDoc('document1.docx')
    loopParse()
    # moonparse()

# See PyCharm help at https://www.jetbrains.com/help/pycharm/
