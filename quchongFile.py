from docx import Document
import re


def remove_duplicate_lines_and_replace_text(doc_path):
    doc = Document(doc_path)

    seen = set()
    paragraphs_to_keep = []

    for para in doc.paragraphs:
        # 替换指定的字符串
        modified_text = para.text.replace('“未找到标题”中提到：“', '“')

        # 分割段落中的每一项
        items = re.split(r'\n|\r', modified_text)
        unique_items = []

        # 检测标题行，保留所有标题行
        is_title = bool(re.match(r'^金句：', para.text.strip()))

        for item in items:
            # 清理每一项，移除数字和点
            cleaned_item = re.sub(r'^\d+\.\s*', '', item.strip())
            if is_title or (cleaned_item and cleaned_item not in seen):
                seen.add(cleaned_item)
                unique_items.append(item.strip())

        if unique_items:
            # 将处理过的独立项重新组合成一个段落
            paragraphs_to_keep.append('\n'.join(unique_items))

    # 清空文档中现有的段落
    for para in list(doc.paragraphs):
        p = para._element
        p.getparent().remove(p)
        p._p = p._element = None

    # 将唯一的段落重新加入到文档
    for line in paragraphs_to_keep:
        doc.add_paragraph(line)

    # 保存修改后的文档
    new_doc_path = doc_path.replace(".docx", "_modified.docx")
    doc.save(new_doc_path)
    return new_doc_path


def main():
    doc_path = 'word_yanbojunjinju.docx'  # 修改为你的文件路径
    result_path = remove_duplicate_lines_and_replace_text(doc_path)
    print("Modified document saved to:", result_path)


if __name__ == "__main__":
    main()
