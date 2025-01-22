import re
from docx import Document
import re

def process_text(content):
    # 使用列表来维护顺序，并使用集合进行去重
    jingles = []
    contra_views = []
    hooks = []

    seen_jingles = set()
    seen_contra_views = set()
    seen_hooks = set()

    category_mapping = {
        "1. 金句": (jingles, seen_jingles),
        "2. 反共识观点": (contra_views, seen_contra_views),
        "3. 钩子语句": (hooks, seen_hooks)
    }

    current_category = None
    lines = content.split("\n")

    for line in lines:
        line = line.strip()
        if line in category_mapping:
            current_category = category_mapping[line]
        elif re.match(r"^\d+\.", line):
            if current_category:
                # 找到第一个点后的位置，并提取从该位置开始的句子
                index = line.find('.') + 1
                sentence = line[index:].strip().strip('"')
                if sentence not in current_category[1]:
                    current_category[0].append(sentence)
                    current_category[1].add(sentence)

    return jingles, contra_views, hooks

def read_word_file(file_path):
    try:
        document = Document(file_path)
        content = []
        for para in document.paragraphs:
            content.append(para.text)
        return "\n".join(content)
    except Exception as e:
        print(f"Error reading file {file_path}: {e}")
        return ""

def write_word_file(jingles, contra_views, hooks, output_path):
    try:
        doc = Document()
        doc.add_paragraph('合并后的内容', style='Title')

        if jingles:
            doc.add_paragraph('金句:', style='Heading 2')
            for jingle in jingles:
                doc.add_paragraph(f'- "{jingle}"', style='Body Text')

        if contra_views:
            doc.add_paragraph('反共识观点:', style='Heading 2')
            for view in contra_views:
                doc.add_paragraph(f'- "{view}"', style='Body Text')

        if hooks:
            doc.add_paragraph('钩子语句:', style='Heading 2')
            for hook in hooks:
                doc.add_paragraph(f'- "{hook}"', style='Body Text')

        doc.save(output_path)
        print(f"Document saved successfully to {output_path}.")
    except Exception as e:
        print(f"An error occurred while writing the file {output_path}: {e}")


if __name__ == '__main__':
    file_path = 'helaoshi1jinju_modified.docx'  # Update this path to your actual file path
    output_path = 'processed_output.docx'  # Update this path to your desired output file path
    content = read_word_file(file_path)
    if content:
        jingles, contra_views, hooks = process_text(content)
        write_word_file(jingles, contra_views, hooks, output_path)
    else:
        print("Failed to read content. Please check the file path and permissions.")
