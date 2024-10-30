import os
import shutil
from os import replace

from docx.shared import Pt
from openai import OpenAI
from docx import Document

import threading
import datetime
import quchongFile

client = OpenAI(
    api_key="aecefa8c98656304dd5f26a3c6656bab.ne82Z5rap89I8H3d",
    base_url="https://open.bigmodel.cn/api/paas/v4/"
)

input_doc_name = "part7.docx"
output_dir = 'output_documents'
output_dir_new = 'output_documents_new'

#每10篇文章聚合成一个
combine_folder = 'output_documents_new_combine'

#4个大的output_documents_new_combine聚合成一个
combine_folder_big = 'output_documents_big_combine'

#最终生成一个合订本
finally_combine_jinju_folder = 'output_documents_total_jinju_combine'
jinju_md_file_path = 'jinju.md'
group_size = 5
big_group_size = 200


def  init():
    # 如果输出目录存在，先清空目录
    if os.path.exists(output_dir):
        shutil.rmtree(output_dir)
    # 重新创建输出目录
    os.makedirs(output_dir)

    if os.path.exists(output_dir_new):
        shutil.rmtree(output_dir_new)
        # 重新创建输出目录
    os.makedirs(output_dir_new)

    if os.path.exists(combine_folder):
        shutil.rmtree(combine_folder)
        # 重新创建输出目录
    os.makedirs(combine_folder)

    if os.path.exists(combine_folder_big):
        shutil.rmtree(combine_folder_big)
        # 重新创建输出目录
    os.makedirs(combine_folder_big)

    if os.path.exists(finally_combine_jinju_folder):
        shutil.rmtree(finally_combine_jinju_folder)
        # 重新创建输出目录
    os.makedirs(finally_combine_jinju_folder)

def splitDoc():
    # 读取输入的Word文档
    input_doc = Document(input_doc_name)

    # 将文档中的所有段落拼接成一个完整的字符串
    full_text = "\n".join([para.text for para in input_doc.paragraphs])

    # 使用 "https://www.douyin.com/video/" 作为分隔符进行文本分割
    sections = full_text.split("https://www.douyin.com/video/")

    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)

    # 遍历每个分割后的部分，从第二部分开始，因为第一部分是分隔前的内容
    for i, section in enumerate(sections[1:], start=1):
        # 跳过包含特定短语的部分
        if "未找到视频ASR文本" in section:
            print(f"跳过包含 '## 未找到视频ASR文本。' 的段落，视频ID: {section.split()[0]}")
            continue  # Skip this section

        if section.strip():  # 忽略空部分
            doc = Document()
            # 恢复视频链接并提取其余部分
            section_parts = section.split("\n", 1)  # 按第一行分割
            video_id = section_parts[0].strip()
            video_link = "https://www.douyin.com/video/" + video_id
            remaining_text = section_parts[1].strip() if len(section_parts) > 1 else ""

            # 找到点赞数部分
            likes_split = remaining_text.split("# 标题:", 1)
            likes_section = likes_split[0].strip()  # 提取点赞数部分
            if likes_section.startswith("# 点赞数:"):
                likes_section = likes_section[len("# 点赞数:"):].strip()

            # 添加点赞数和视频链接到文档
            doc.add_paragraph(video_link)
            doc.add_paragraph(likes_section)

            # 添加标题和后续内容
            if len(likes_split) > 1:
                title_and_content = "# 标题:" + likes_split[1].strip()
                doc.add_paragraph(title_and_content)

            # 定义输出路径
            output_path = os.path.join(output_dir, f'{input_doc_name.replace(".docx", "")}_{i}.docx')
            doc.save(output_path)
            print(f"文档已保存: {output_path}")

    print(f"文档拆分完成，共生成 {len(sections) - 1} 个文件，保存在目录: {output_dir}")


def process_files(file_list, thread_name):
    for filepath in file_list:
        print(f"{thread_name} - 正在处理文件: {filepath}")
        zhipuparse(filepath)

def loopParse():
    output_dir = 'output_documents'
    if os.path.exists(output_dir) and os.path.isdir(output_dir):
        print(f"'{output_dir}' 文件夹存在，正在遍历其中的文件...")
        docx_files = [os.path.join(output_dir, filename) for filename in os.listdir(output_dir)
                      if filename.endswith('.docx') and not filename.startswith('~$')]

        # 将文件列表分成五部分
        fifth = len(docx_files) // 5
        files_part1 = docx_files[:fifth]
        files_part2 = docx_files[fifth:2 * fifth]
        files_part3 = docx_files[2 * fifth:3 * fifth]
        files_part4 = docx_files[3 * fifth:4 * fifth]
        files_part5 = docx_files[4 * fifth:]

        # 创建五个线程分别处理五部分文件
        thread1 = threading.Thread(target=process_files, args=(files_part1, "线程1"))
        thread2 = threading.Thread(target=process_files, args=(files_part2, "线程2"))
        thread3 = threading.Thread(target=process_files, args=(files_part3, "线程3"))
        thread4 = threading.Thread(target=process_files, args=(files_part4, "线程4"))
        thread5 = threading.Thread(target=process_files, args=(files_part5, "线程5"))

        # 启动线程
        thread1.start()
        thread2.start()
        thread3.start()
        thread4.start()
        thread5.start()

        # 等待所有线程完成
        thread1.join()
        thread2.join()
        thread3.join()
        thread4.join()
        thread5.join()

        print("所有文件处理完成。")

    else:
        print(f"'{output_dir}' 文件夹不存在。")


def combine_word_documents():
    # 创建输出文件夹如果它不存在
    if not os.path.exists(combine_folder):
        os.makedirs(combine_folder)

    # 获取所有Word文档
    word_files = [f for f in os.listdir(output_dir_new) if f.endswith('.docx')]
    word_files.sort()  # 可选：排序文件

    # 确定需要创建的文档数量
    num_docs = len(word_files) // group_size + (1 if len(word_files) % group_size else 0)

    # 合并文档
    for i in range(num_docs):
        # 创建新文档
        merged_document = Document()
        start_index = i * group_size
        end_index = start_index + group_size

        # 添加每个文件的内容到新文档
        for j in range(start_index, min(end_index, len(word_files))):
            doc_path = os.path.join(output_dir_new, word_files[j])
            sub_doc = Document(doc_path)
            for para in sub_doc.paragraphs:  # 正确的添加段落内容
                merged_document.add_paragraph(para.text)

            # 添加页面断裂（如果需要）
            if j < end_index - 1:  # 不在最后一个文档后添加分页符
                merged_document.add_page_break()

        # 保存合并后的文档
        output_path = os.path.join(combine_folder, f'merged_document_{i + 1}.docx')
        merged_document.save(output_path)
        print(f'Saved merged document {output_path}')


def delete_files_in_folder(folder_path):
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path):
            os.remove(file_path)



def combine_word_documents_for_big():
    # 创建输出文件夹如果它不存在
    if not os.path.exists(combine_folder_big):
        os.makedirs(combine_folder_big)

    # 获取所有Word文档
    word_files = [f for f in os.listdir(combine_folder) if f.endswith('.docx')]
    word_files.sort()  # 可选：排序文件

    # 确定需要创建的文档数量
    num_docs = len(word_files) // big_group_size + (1 if len(word_files) % big_group_size else 0)

    # 合并文档
    for i in range(num_docs):
        # 创建新文档
        merged_document = Document()
        start_index = i * big_group_size
        end_index = start_index + big_group_size

        # 添加每个文件的内容到新文档
        for j in range(start_index, min(end_index, len(word_files))):
            doc_path = os.path.join(combine_folder, word_files[j])
            sub_doc = Document(doc_path)
            for para in sub_doc.paragraphs:  # 正确的添加段落内容
                merged_document.add_paragraph(para.text)

            # 添加页面断裂（如果需要）
            if j < end_index - 1:  # 不在最后一个文档后添加分页符
                merged_document.add_page_break()

        # 保存合并后的文档
        output_path = os.path.join(combine_folder_big, f'merged_document_{input_doc_name}')
        merged_document.save(output_path)
        print(f'Saved merged document {output_path}')

#调用智谱清言 解析
def zhipuparse(filepath):
    try:
        # 创建新文件夹
        os.makedirs(output_dir_new, exist_ok=True)  # 确保文件夹存在
        # 加载文件并调用API

        # 读取输入的Word文档
        input_doc = Document(filepath)

        # 将文档中的所有段落拼接成一个完整的字符串
        full_text = "\n".join([para.text for para in input_doc.paragraphs])

        sections = full_text.split("视频ASR文本:")

        # 第一步：检查是否存在第二部分
        if len(sections) < 2 or not sections[1].strip():
            sections = full_text.split("视频ASR文本：")  # 使用全角冒号的分隔符

        # 第二步：如果仍然没有找到有效部分，尝试使用 "视频或图片OCR文本:" 进行分割
        if len(sections) < 2 or not sections[1].strip():
            sections = full_text.split("视频或图片OCR文本:")

        # 第三步：最后一次检查是否存在有效文本
        if len(sections) < 2 or not sections[1].strip():
            sections = full_text.split("视频或图片OCR文本：")

        if len(sections) < 2 or not sections[1].strip():
            print(f"文件 {filepath} 中没有找到有效的 '视频ASR文本:' 或 '视频或图片OCR文本:' 部分，跳过该文件的处理")


        asr_text = sections[1].strip()  # 如果找到合适的分割文本，获取并去除首尾空格

        completion = client.chat.completions.create(
            model="glm-4",
            messages=[
                {"role": "system", "content": "你是一位中文专家，标点符号专家"},
                {"role": "user",
                 "content": asr_text + "  请针对以上这段中的“## 视频ASR文本：”部分进行标点符号的补全，如果有错别字，请一并修订。"
                                       "直接输出结果，不需要给出修改提示，例如这样的提示\"修改后的文本已经补全了标点符号，并修正了一些语序上的小问题，以提高语句的通顺性和可读性。\"。"
                                       "不要输出修改了哪些内容的提示，例如'（注：以上文本中，“g l j ”应为不规范的缩写或打字错误，但因为没有上下文信息，无法确定具体应为何词或短语，故不做修改。"
                                       "请务必注意不要保留修改提示语句，例如“（注：以上文本中，错别字和不通顺的语句已经根据上下文进行了修改，并补全了标点符号。 还有类似这种“由于原文中包含很多方言和口语表达，某些地方可能理解有误，未能完全反映原意，还请谅解。””"
                                       "再比如这种提示词，也不要保留”（注：文本中已补全标点符号，并修正了部分语序和用词，以提高语句的通顺性和可读性。“"
                                      }

            ],
            top_p=0.7,
            temperature=0.9
        )
        print(completion.choices[0].message)
        result = str(completion.choices[0].message)
        result =  (((result.replace("', refusal=None, role='assistant', function_call=None, tool_calls=None)","")
                   .replace("', refusal=None, role='assistant', function_call=None, tool_calls=None)","")
                   .replace("。）", "。")
                   .replace("ChatCompletionMessage(content='","")).replace("\\n","\n")
                   .replace(", role='assistant', function_call=None, tool_calls=None)",""))
                   .replace("你是一位中文专家，标点符号专家。",""))
        # 加载文档document_modified_2
        input_doc = Document(filepath)

        # 遍历文档中的段落
        for paragraph in input_doc.paragraphs:
            if "视频ASR文本:" in paragraph.text or "视频ASR文本：" in paragraph.text \
                    or "视频或图片OCR文本:" in paragraph.text or "视频或图片OCR文本：" in paragraph.text:
                # 使用多种分隔符进行分割
                parts = paragraph.text.split("视频ASR文本:") if "视频ASR文本:" in paragraph.text else \
                    paragraph.text.split("视频ASR文本：") if "视频ASR文本：" in paragraph.text else \
                        paragraph.text.split("视频或图片OCR文本:") if "视频或图片OCR文本:" in paragraph.text else \
                            paragraph.text.split("视频或图片OCR文本：")

                # 替换文本中的内容为API返回的结果
                if len(parts) > 1:
                    paragraph.text = parts[0] + result.replace("ChatCompletionMessage(content='## ", "")
                    # 输出替换后的段落内容，检查标点符号是否存在
                    # print(f"替换后的段落内容: {paragraph.text}")
        output_path = os.path.join(output_dir_new, os.path.basename(filepath))
        input_doc.save(output_path)
        print(f"保存路径：{output_path}")

    except Exception as e:
        print(f"处理文件 {filepath} 时发生错误: {e}")
        full_file_path = output_dir_new + "/" + os.path.basename(filepath)  # 获取文件名
        if os.path.exists(full_file_path):
            os.remove(full_file_path)
            print(f"文件 {full_file_path} 已被删除。")
        else:
            print(f"文件 {full_file_path} 不存在。")


# 提取金句
def tiqujinju(filename):
    if filename.endswith('.docx') and not filename.startswith("$"):  # 检查文件扩展名是否为.docx
        file_path = filename
        try:
            input_doc = Document(file_path)
            # 将文档中的所有段落拼接成一个完整的字符串
            full_text = "\n".join([para.text for para in input_doc.paragraphs])

            completion = client.chat.completions.create(
                model="glm-4",
                messages=[
                    {"role": "system", "content": "你是一位自媒体领域的大V"},
                    {"role": "user",
                     "content": full_text + "  请提取分析其中的金句5-10句，反共识观点5-10句，以及最明显的钩子语句5-10句。请直接输出结果。所谓钩子语句是指："
                                            "在自媒体领域的核心原则是给用户一个私信你的理由，通过提供独特价值或个性化帮助，激发用户的兴趣和行动。它不仅要吸引注意力，还要让读者感到如果不联系你，就会错失某种重要机会或资源。"
                                            "注意你需要直接给出相应的钩子，无需做任何解释，特别是无需解释为什么你认为它是钩子。"
                                           "请不要输出类似'以上提取的金句、反共识观点和钩子语句均未涉及不安全或敏感内容。'这样的提示语"}

                ],
                top_p=0.7,
                temperature=0.9
            )
            jinjiu_content = str(completion.choices[0].message)
            print(jinjiu_content)

            # 文件路径
            output_file_path = os.path.join(finally_combine_jinju_folder,
                                            input_doc_name.replace(".docx","") + "jinju.docx")

            # 如果文件不存在，则创建一个新文档
            if not os.path.exists(output_file_path):
                doc = Document()
            else:
                doc = Document(output_file_path)

            # 替换文本中的特定内容
            modified_content = (jinjiu_content \
                .replace("', refusal=None, role='assistant', function_call=None, tool_calls=None)", "") \
                .replace("ChatCompletionMessage(content='", "").replace("', role='assistant', function_call=None, tool_calls=None)","")
                .replace("你是一位中文专家，标点符号专家，以下是针对你提供文本的标点符号补全和错别字的修订：",""))
            # 将内容按行添加到文档
            for line in modified_content.splitlines():
                p = doc.add_paragraph(line)
                p.style.font.size = Pt(12)  # 例如设置字体大小为12pt
            # 保存文档
            doc.save(output_file_path)
            quchongFile.remove_duplicate_lines_and_replace_text(output_file_path)

        except Exception as e:
            print(f"处理文件 {file_path} 时发生错误: {e}")
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"文件 {file_path} 已被删除。")
            else:
                print(f"文件 {file_path} 不存在。")

#处理分组文件
def process_files_jinju(file_list, thread_name):
    for filepath in file_list:
        print(f"{thread_name} - 正在处理文件: {filepath}")
        tiqujinju(filepath)

def loopParsetiqujinju():
    # 遍历指定文件夹
    # 创建输出文件夹如果它不存在
    if not os.path.exists(finally_combine_jinju_folder):
        os.makedirs(finally_combine_jinju_folder)
    else:
        delete_files_in_folder(finally_combine_jinju_folder)


    if os.path.exists(combine_folder) and os.path.isdir(combine_folder):
        print(f"'{combine_folder}' 文件夹存在，正在遍历其中的文件...")
        docx_files = [os.path.join(combine_folder, filename) for filename in os.listdir(combine_folder)
                      if filename.endswith('.docx') and not filename.startswith('~$')]
        # 将文件列表分成五部分
        fifth = len(docx_files) // 5
        files_part1 = docx_files[:fifth]
        files_part2 = docx_files[fifth:2 * fifth]
        files_part3 = docx_files[2 * fifth:3 * fifth]
        files_part4 = docx_files[3 * fifth:4 * fifth]
        files_part5 = docx_files[4 * fifth:]

        # 创建五个线程分别处理五部分文件
        thread1 = threading.Thread(target=process_files_jinju, args=(files_part1, "线程1"))
        thread2 = threading.Thread(target=process_files_jinju, args=(files_part2, "线程2"))
        thread3 = threading.Thread(target=process_files_jinju, args=(files_part3, "线程3"))
        thread4 = threading.Thread(target=process_files_jinju, args=(files_part4, "线程3"))
        thread5 = threading.Thread(target=process_files_jinju, args=(files_part5, "线程3"))

        # 启动线程
        thread1.start()
        thread2.start()
        thread3.start()
        thread4.start()
        thread5.start()

        # 等待所有线程完成
        thread1.join()
        thread2.join()
        thread3.join()
        thread4.join()
        thread5.join()
        print("所有文件处理完成。")


if __name__ == '__main__':
    # 记录开始时间
   start_time = datetime.datetime.now()
   print(f"开始时间: {start_time}")
   init()
   splitDoc()
   loopParse()
   combine_word_documents()
   combine_word_documents_for_big()
   loopParsetiqujinju()
   end_time = datetime.datetime.now()
   print(f"结束时间: {end_time}")
   # 计算并打印运行时间
   time_taken = end_time - start_time
   print(f"运行时间: {time_taken}")
