import os
import time
from pathlib import Path
from docx import Document
from openai import OpenAI

# 初始化 API 客户端（需要确保 API Key 正确）
client = OpenAI(
    api_key="sk-gyqGihWcBDXjapcnDnOuDeLDDTgQ2oNmM5DUBmd5urZhwpl9",
    base_url="https://api.moonshot.cn/v1",
)

input_doc_path = "document1.docx"
output_dir = 'output_documents'


def splitDoc():
    # 读取输入的Word文档
    input_doc = Document(input_doc_path)

    # 将文档中的所有段落拼接成一个完整的字符串
    full_text = "\n".join([para.text for para in input_doc.paragraphs])

    # 使用 "https://www.douyin.com/video/" 作为分隔符进行文本分割
    sections = full_text.split("https://www.douyin.com/video/")

    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)

    # 遍历每个分割后的部分，从第二部分开始，因为第一部分是分隔前的内容
    for i, section in enumerate(sections[1:], start=1):
        if section.strip():  # 忽略空部分
            doc = Document()
            # 恢复视频链接并提取其余部分
            section_parts = section.split("\n", 1)  # 按第一行分割
            video_id = section_parts[0].strip()
            video_link = "https://www.douyin.com/video/" + video_id
            remaining_text = section_parts[1].strip() if len(section_parts) > 1 else ""

            # 找到点赞数部分
            likes_split = remaining_text.split("# 标题:", 1)
            likes_section = likes_split[0].strip()  # 提取点赞数部分
            if likes_section.startswith("# 点赞数:"):
                likes_section = likes_section[len("# 点赞数:"):].strip()

            # 添加点赞数和视频链接到文档
            doc.add_paragraph(video_link)
            doc.add_paragraph(likes_section)

            # 添加标题和后续内容
            if len(likes_split) > 1:
                title_and_content = "# 标题:" + likes_split[1].strip()
                doc.add_paragraph(title_and_content)

            # 定义输出路径
            output_path = os.path.join(output_dir, f'document_{i}.docx')
            doc.save(output_path)
            print(f"文档已保存: {output_path}")

    print(f"文档拆分完成，共生成 {len(sections) - 1} 个文件，保存在目录: {output_dir}")

def loopParse():
    output_dir = 'output_documents'
    if os.path.exists(output_dir) and os.path.isdir(output_dir):
        print(f"'{output_dir}' 文件夹存在，正在遍历其中的文件...")
        for filename in os.listdir(output_dir):
            if filename.endswith('.docx') and not filename.startswith('~$'):
                filepath = os.path.join(output_dir, filename)
                print(f"正在处理文件: {filepath}")
                moonparse(filepath)
    else:
        print(f"'{output_dir}' 文件夹不存在。")

def moonparse(filepath):

    try:
        # 创建新文件夹
        output_dir_new = 'output_documents_new'
        os.makedirs(output_dir_new, exist_ok=True)  # 确保文件夹存在

        # 设置新的文件路径
        filename = os.path.basename(filepath)  # 获取文件名
        output_path = os.path.join(output_dir_new, filename)

        # 加载文件并调用API
        file_object = client.files.create(file=Path(filepath), purpose="file-extract")
        file_content = client.files.content(file_id=file_object.id).text

        # 执行API请求并获取处理结果
        messages = [
            {"role": "system", "content": "你是 Kimi，由 Moonshot AI 提供的人工智能助手..."},
            {"role": "system", "content": file_content},
            {"role": "user", "content": "仅对文档中的“## 视频ASR文本：”部分进行标点符号的补全。直接输出结果，不需要给出修改提示"}
        ]

        # 然后调用 chat-completion, 获取 Kimi 的回答
        completion = client.chat.completions.create(
            model="moonshot-v1-128k",
            messages=messages,
            temperature=0.3,
        )

        result = str(completion.choices[0].message)
        result = (result.replace("ChatCompletionMessage(content='您好，", "")
                  .replace("（文档结束）', refusal=None, role='assistant', function_call=None, tool_calls=None)", "")
                  .replace("', refusal=None, role='assistant', function_call=None, tool_calls=None)", "")
                  .replace("refusal=None, role='assistant', function_call=None, tool_calls=None)", "")
                  .replace("ChatCompletionMessage(content=\'", ""))

        # 加载文档document_modified_2
        input_doc = Document(filepath)

        # 遍历文档中的段落
        for paragraph in input_doc.paragraphs:
            if "## 视频ASR文本:" in paragraph.text:
                # 分割文本以保留 "## 视频ASR文本:" 前的内容
                parts = paragraph.text.split("## 视频ASR文本:")
                # 重新组合并替换后面的内容为 "暂无文档说明"
                paragraph.text = parts[0] + result.replace("ChatCompletionMessage(content='## ", "")

        # 保存修改后的文档
        output_path = 'output_documents_new/' + filename
        input_doc.save(output_path)

    except Exception as e:
        print(f"处理文件 {filepath} 时发生错误: {e}")


if __name__ == '__main__':
    # splitDoc()
    loopParse()
